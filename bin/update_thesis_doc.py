from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
THESIS_PATH = ROOT / "docs" / "待提交" / "论文.docx"
BACKUP_PATH = ROOT / "docs" / "待提交" / "论文_修改前备份.docx"
TEMP_OUTPUT_PATH = ROOT / "docs" / "待提交" / "论文_自动更新临时稿.docx"

ARCH_IMAGE = ROOT / "docs" / "assets" / "figures" / "system-architecture.png"
SCREEN_HOME = ROOT / "docs" / "assets" / "screenshots" / "home-dashboard-viewport-top.png"
SCREEN_LIST = ROOT / "docs" / "assets" / "screenshots" / "drivinglog-list-viewport.png"
SCREEN_FORECAST = ROOT / "docs" / "assets" / "screenshots" / "forecast-workbench-viewport.png"
SCREEN_NASA = ROOT / "docs" / "assets" / "screenshots" / "forecast-nasa-viewport.png"
NASA_CURVE = ROOT / "artifacts" / "battery_life" / "figures" / "nasa_capacity_degradation.png"


def backup_source() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(THESIS_PATH, BACKUP_PATH)


def paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> Paragraph:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def insert_picture_after(paragraph: Paragraph, image_path: Path, width: float = 5.8) -> Paragraph:
    pic_para = paragraph_after(paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    return pic_para


def insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Table:
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.2))
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)
    return table


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"未找到段落: {text}")


def find_paragraph_contains(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text in paragraph.text.strip():
            return paragraph
    raise ValueError(f"未找到包含指定内容的段落: {text}")


def append_reference_if_missing(doc: Document, line: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == line.strip():
            return
    ref_heading = find_paragraph(doc, "参考文献")
    last_ref = ref_heading
    seen_ref_heading = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "参考文献":
            seen_ref_heading = True
            continue
        if not seen_ref_heading:
            continue
        if paragraph.text.strip() == "致谢":
            break
        if paragraph.text.strip():
            last_ref = paragraph
    paragraph_after(last_ref, line, style=last_ref.style.name if last_ref.style else None)


def build_metric_table(anchor: Paragraph) -> Paragraph:
    caption = paragraph_after(anchor, "表4-1 业务预测与 NASA 实验指标汇总")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = insert_table_after(caption, 7, 7)
    headers = ["实验模块", "任务", "样本数", "MAE", "RMSE", "R²", "说明"]
    rows = [
        ["业务主模型", "能耗预测", "21507", "0.202", "1.192", "0.514", "系统默认展示结果"],
        ["业务主模型", "电量寿命代理预测", "21507", "0.000", "0.000", "1.000", "标签规则性较强，结果不宜过度解读"],
        ["深度学习对比", "能耗预测", "21467", "0.389", "1.161", "0.638", "用于方法对比，不作为默认输出"],
        ["深度学习对比", "电量寿命代理预测", "21467", "0.034", "0.069", "0.000", "在该任务上稳定性弱于主模型"],
        ["NASA 实验", "SOH 预测", "5308", "0.059", "0.115", "0.600", "27 节电池训练，7 节电池测试"],
        ["NASA 实验", "RUL 预测", "5308", "13.339", "24.996", "0.450", "组间留出验证，体现真实寿命建模能力"],
    ]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = value
    return caption


def build_test_table(anchor: Paragraph) -> Paragraph:
    caption = paragraph_after(anchor, "表5-1 系统测试用例")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = insert_table_after(caption, 7, 4)
    headers = ["编号", "测试项", "用例说明", "预期结果"]
    rows = [
        ["TC-01", "登录态清理", "验证登录、退出登录和鉴权失败时的本地会话清理逻辑", "旧调试状态被清理，页面可重新登录"],
        ["TC-02", "后端启动脚本", "执行 `start_backend.ps1`，检查 `.env` 读取和 WSGI 启动流程", "后端能够稳定启动并监听设定端口"],
        ["TC-03", "环境变量读取", "验证前后端 `.env.example` 的端口与 API 地址配置", "前后端地址保持一致"],
        ["TC-04", "车型知识采集", "调用车型知识采集接口并检查去重、回退与入库结果", "接口返回成功，知识库记录正确生成"],
        ["TC-05", "预测工作台接口", "检查模型指标、场景预测和对比接口返回结构", "前端可正常展示业务预测结果"],
        ["TC-06", "NASA 实验接口", "调用 `nasaExperiment` 接口并读取图像与指标清单", "返回 `code=0` 且实验图像可访问"],
    ]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = value
    return caption


def main() -> None:
    backup_source()
    doc = Document(THESIS_PATH)

    set_paragraph_text(
        find_paragraph_contains(doc, "本课题围绕电动汽车行车日志分析系统展开设计与实现"),
        "本文面向电动汽车运行数据分析与可视化需求，设计并实现了一套电动汽车行车日志分析系统。系统以 Django 为后端、Vue 3 为前端，完成了行车日志管理、统计看板、车型知识采集、预测工作台和管理员后台等核心功能。在数据分析方面，系统以结构化行车日志为基础构建机器学习主模型，并补充深度学习对比实验，用于比较不同方法在业务数据上的表现。为增强论文的研究深度，本文进一步引入 NASA PCoE Battery Aging 公开数据集，完成电池 SOH 和 RUL 预测实验，形成与业务系统相互独立但可相互支撑的实验链路。工程实现上，项目补充了一键启动脚本、统一环境变量、前端会话清理和基础自动化测试，提升了系统的可运行性和可演示性。实践表明，该系统能够较完整地支撑电动汽车行车日志的管理、分析与预测展示任务，也具备一定的研究扩展空间。",
    )
    set_paragraph_text(
        find_paragraph(doc, "关键词：电动汽车；行车日志分析；分布式处理；数据可视化；智能分析模型"),
        "关键词：电动汽车；行车日志分析；Django；Vue 3；机器学习；SOH/RUL预测",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "This project focuses on the design and implementation of an electric vehicle"),
        "This thesis designs and implements an electric vehicle driving-log analysis system oriented to data management, visual analytics and intelligent prediction. The project uses Django as the backend and Vue 3 as the frontend, and integrates driving-log management, dashboard visualization, vehicle knowledge crawling, a prediction workbench and an administrator console. For intelligent analysis, a machine-learning pipeline is used as the default business prediction path, while a deep-learning comparison pipeline is added for method comparison. In addition, the NASA PCoE Battery Aging dataset is introduced to conduct SOH and RUL experiments, which strengthens the research depth of the thesis while remaining independent from the business module. The final project also includes startup scripts, unified environment-variable configuration and basic automated tests, improving reproducibility and demonstration stability.",
    )
    set_paragraph_text(
        find_paragraph(doc, "Keywords: Electric vehicle; Driving log analysis; Distributed processing; Data visualization; Intelligent analysis model"),
        "Keywords: Electric vehicle; Driving log analysis; Django; Vue 3; Machine learning; SOH/RUL prediction",
    )

    set_paragraph_text(
        find_paragraph_contains(doc, "国内对电动汽车数据处理与分析的研究随新能源产业发展不断深入"),
        "国内相关研究主要围绕新能源汽车运行数据采集、远程监控、报文分析和平台化展示展开，研究重点多放在运行状态监测、异常预警和基础可视化建设上。已有工作表明，基于大数据平台对电动汽车报文和运行日志进行汇总分析，能够为车辆状态感知和运维管理提供支持[12]。不过，从当前公开成果来看，许多系统仍以管理展示为主，对预测模型、实验验证和系统化工程收口的结合还不够充分。",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "国外新能源汽车产业起步较早"),
        "国外研究在电动汽车能耗预测和锂电池寿命预测两个方向上发展较快。在能耗预测方面，Pan 等人提出将短行程分段与深度学习结合，用于真实工况下的电动汽车能耗预测[16]；Ko 等人验证了 GPS 轨迹数据与深度学习结合进行车辆能耗监测的可行性[17]；Pamuła 等人进一步将深度学习用于电动公交行程能耗预测[18]。在电池寿命研究方面，Li 等人构建了同时面向 SOH 与 RUL 的端到端神经网络框架[19]，Zhu 等人提出注意力 CNN-BiLSTM 模型[20]，Zhou 等人使用 TCN 开展电池状态监测与剩余寿命预测[21]，Ren 等人与 Wang 等人则分别从 CNN-LSTM 组合建模和迁移学习角度改进了寿命预测性能[22][23]。",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "综合来看，国内外均高度重视电动汽车行车日志数据的价值挖掘"),
        "综合来看，现有研究已经证明了电动汽车运行数据和电池退化数据具有较高的分析价值，但仍存在两个明显问题：一是工程系统与算法实验常常分离，论文中有模型但系统里落不下去；二是业务系统中的代理标签预测与公开真实电池寿命实验缺少清晰边界。基于此，本文尝试把行车日志分析系统的工程实现与 NASA 电池 SOH/RUL 实验放在同一篇论文中完成，但在数据来源、实验目标和结论解释上保持明确区分，从而兼顾系统可用性与研究完整性。",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "本文旨在设计并实现一套功能完善、运行稳定的电动汽车行车日志分析系统"),
        "本文的研究目标是设计并实现一套可运行、可展示、可验证的电动汽车行车日志分析系统。一方面，系统面向业务场景完成日志管理、统计分析、车型知识采集、预测工作台和后台管理等功能；另一方面，围绕智能分析模块，构建机器学习主模型、深度学习对比实验以及 NASA 公开电池老化数据集上的 SOH/RUL 实验，形成较为完整的系统与实验闭环。同时，通过统一环境变量、一键启动脚本和自动化测试，提高项目在答辩场景下的稳定性与复现性。",
    )

    set_paragraph_text(
        find_paragraph_contains(doc, "本课题围绕电动汽车行车日志分析系统展开开发"),
        "本课题围绕电动汽车行车日志分析系统展开实现，当前答辩版本以前后端分离 Web 架构为主体，使用 Django 提供后端接口与训练命令，使用 Vue 3 与 ECharts 构建管理端界面，并结合机器学习、深度学习对比实验和 NASA 电池 SOH/RUL 实验完成智能分析模块。Hadoop 与 Hive 在论文中仍作为面向海量数据扩展的设计基础保留，用于说明系统在大规模日志场景下的可扩展方向。",
    )

    p31 = find_paragraph_contains(doc, "本课题设计的电动汽车行车日志分析系统采用分层式分布式架构")
    set_paragraph_text(
        p31,
        "本系统采用分层式架构设计，按照“数据来源层、数据治理层、数据存储层、算法分析层、服务接口层、前端展示层”进行组织。数据来源层同时接收业务行车日志、车型知识公开网页和 NASA 电池老化公开数据；数据治理层负责清洗、字段规范化与特征构造；数据存储层以 MySQL 和 `artifacts/` 训练产物目录为主体，并配合 Redis 完成缓存与会话管理；算法分析层由机器学习主模型、深度学习对比链路和 NASA SOH/RUL 实验组成；服务接口层基于 Django 实现统一访问入口；前端展示层基于 Vue 3 构建首页、日志列表、预测工作台和 NASA 实验页，形成从数据到展示的完整闭环。",
    )
    image_anchor = insert_picture_after(p31, ARCH_IMAGE, width=6.2)
    set_paragraph_text(find_paragraph(doc, "图 3-1 模块图"), "图3-1 系统总体架构图", align=WD_ALIGN_PARAGRAPH.CENTER)

    set_paragraph_text(
        find_paragraph_contains(doc, "数据采集是电动汽车行车日志分析系统运行的首要环节"),
        "在当前系统中，数据采集分为三类来源：第一类是业务行车日志数据，主要通过现有 Excel 数据和数据库导入完成；第二类是车型知识数据，通过公开网页采集补全车辆品牌、电池类型、续航等说明性信息；第三类是 NASA PCoE Battery Aging 公开数据集，用于电池健康状态与剩余寿命实验。这样的设计使系统既能支撑业务演示，也能保留研究实验所需的真实公共数据来源。",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "数据预处理是指对爬取的原始数据集进行一系列的操作"),
        "数据预处理的目标是把来源不同、质量不一的原始数据转换为可用于建模和展示的结构化样本。针对行车日志数据，系统重点完成缺失值处理、异常值识别、字段标准化、数值归一化和特征构造；针对 NASA 电池数据，则进一步进行循环级样本展开、容量衰减特征提取和训练/测试电池分组。经过这一过程后，业务日志预测链路与 NASA 实验链路都能够形成稳定可复用的数据输入。",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "数据存储是电动汽车行车日志分析系统的底层支撑"),
        "数据存储采用“业务数据与训练产物分离”的方式进行组织。结构化业务数据统一落在 MySQL 中，便于接口查询、后台管理和统计分析；模型、指标、图像和实验清单统一保存在 `artifacts/` 目录，便于前端直接展示和答辩时复用；Redis 负责缓存与会话相关能力。对于更大规模日志场景，论文保留了 HDFS 与 Hive 的扩展设计思路，但在当前答辩交付版本中，核心运行路径已收敛为 MySQL + Django + Vue 3 的稳定组合。",
    )

    analysis_anchor = find_paragraph_contains(doc, "数据分析是电动汽车行车日志分析系统挖掘数据价值、实现智能应用的核心环节")
    set_paragraph_text(
        analysis_anchor,
        "数据分析模块由三条链路构成。第一条链路是业务预测链路，以结构化行车日志为基础，构建能耗与电量寿命代理预测模型，用于系统默认展示；第二条链路是深度学习对比实验，通过序列样本对业务预测任务进行补充性比较；第三条链路是 NASA 公开数据集上的 SOH/RUL 实验，用于验证系统在真实电池寿命建模场景中的可行性。当前版本中，业务主模型在 21507 条样本上的能耗预测取得了 MAE=0.202、RMSE=1.192、R²=0.514 的结果；NASA 实验在 5308 个样本点上的 SOH 预测取得了 MAE=0.059、RMSE=0.115、R²=0.600，在 RUL 预测上取得了 MAE=13.339、RMSE=24.996、R²=0.450，说明系统已具备基础的真实寿命实验能力。",
    )
    metric_caption = build_metric_table(analysis_anchor)
    discussion = paragraph_after(metric_caption, "从结果看，机器学习主模型更适合作为系统默认输出，因为它在当前业务数据条件下更稳定、可解释性更强；深度学习链路在能耗场景中表现出一定潜力，但在电量寿命代理标签任务上并未全面优于主模型。需要说明的是，业务系统中的寿命代理预测与 NASA 真实 SOH/RUL 实验属于两条不同的分析线，结论解释不能混用。")
    curve_pic = insert_picture_after(discussion, NASA_CURVE, width=6.0)
    curve_caption = paragraph_after(curve_pic, "图4-1 NASA 电池容量衰减曲线")
    curve_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_paragraph_text(
        find_paragraph_contains(doc, "数据可视化是电动汽车行车日志分析系统实现数据价值直观呈现"),
        "数据可视化与交互实现是本系统对外展示的核心部分。当前版本以前端后台为主要演示入口，围绕首页总览、行车日志列表、预测工作台和 NASA 实验页组织页面结构，并通过统一主题样式、图表卡片和表单交互提升可读性。与早期模板化界面相比，答辩版本更强调页面层次、预测结果表达和关键业务链路的一致性。",
    )
    set_paragraph_text(
        find_paragraph_contains(doc, "用户模块基于 Vue.js+ECharts 开发"),
        "用户侧模块主要提供注册登录、行车日志浏览、收藏、评论反馈和基础个性化推荐等能力。该部分强调对单条日志与个人使用体验的支撑，适合展示系统在前台交互层面的完整性，也为管理员侧统计分析与知识采集提供业务数据基础。",
    )
    admin_anchor = find_paragraph_contains(doc, "管理员模块基于 Vue.js+ECharts 开发")
    set_paragraph_text(
        admin_anchor,
        "管理员模块是本项目答辩演示的重点页面。其主要包含四类能力：一是首页看板，用于展示日志规模、趋势图和核心入口；二是行车日志列表，用于完成查询、查看、收藏、评论与车型知识采集；三是预测工作台，用于展示机器学习主模型结果、深度学习对比结果及场景模拟；四是 NASA 实验页，用于展示真实电池数据上的 SOH/RUL 指标与图像。通过这些页面，系统可以较完整地体现从业务管理到智能分析再到研究实验验证的全流程实现。",
    )

    anchor = admin_anchor
    for image_path, caption_text in [
        (SCREEN_HOME, "图5-9 管理后台首页"),
        (SCREEN_LIST, "图5-10 行车日志列表界面"),
        (SCREEN_FORECAST, "图5-11 预测工作台界面"),
        (SCREEN_NASA, "图5-12 NASA 实验界面"),
    ]:
        anchor = insert_picture_after(anchor, image_path, width=5.8)
        caption = paragraph_after(anchor, caption_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor = caption

    test_heading = paragraph_after(anchor, "5.3 系统测试与运行验证", style="Heading 2")
    test_intro = paragraph_after(test_heading, "为了保证答辩交付版本的稳定性，系统将测试重点放在登录鉴权、环境变量、启动脚本、关键业务接口与模型实验可用性五个方面。测试既覆盖了工程侧容易出错的启动与配置问题，也覆盖了业务侧最关键的预测与实验接口。")
    test_caption = build_test_table(test_intro)
    paragraph_after(
        test_caption,
        "测试结果表明，前端登录态清理测试与后端启动脚本测试均已通过，环境变量样例文件能够统一约束前后端端口与 API 地址；后端 `manage.py check` 能正常通过；NASA 实验接口在真实运行中返回 `code=0`、`available=True`，并能正确提供 3 张实验图像。这说明当前项目已经具备较好的演示稳定性和基础可复现性。",
    )

    set_paragraph_text(
        find_paragraph_contains(doc, "本课题以电动汽车行车日志分析系统的设计与实现为研究核心"),
        "本文围绕电动汽车行车日志分析系统的设计与实现，完成了一套面向答辩交付的完整系统。项目在工程层面实现了日志管理、可视化展示、车型知识采集、预测工作台、NASA 电池 SOH/RUL 实验页及管理员后台；在算法层面形成了机器学习主模型、深度学习对比实验和真实公开数据集实验三条分析链路；在交付层面补充了统一环境变量、一键启动脚本、登录态清理和基础自动化测试。整体来看，本系统已经具备较好的功能完整性、演示稳定性和研究表达能力。后续工作仍可围绕更真实的车辆采集数据、更完整的深度学习训练环境以及更高覆盖率的自动化测试继续展开。",
    )

    new_refs = [
        "Pan Y, Fang W, Zhang W. Development of an energy consumption prediction model for battery electric vehicles in real-world driving: A combined approach of short-trip segment division and deep learning[J]. Journal of Cleaner Production, 2023, 406: 136742.",
        "Ko K, Lee T, Jeong S. A Deep Learning Method for Monitoring Vehicle Energy Consumption with GPS Data[J]. Sustainability, 2021, 13(20): 11331.",
        "Pamuła T, Pamuła D. Prediction of Electric Buses Energy Consumption from Trip Parameters Using Deep Learning[J]. Energies, 2022, 15(5): 1747.",
        "Li P, Zhang Z, Grosu R, et al. An end-to-end neural network framework for state-of-health estimation and remaining useful life prediction of electric vehicle lithium batteries[J]. Renewable and Sustainable Energy Reviews, 2022, 151: 111843.",
        "Zhu Z, Yang Q, Liu X, Gao D. Attention-based CNN-BiLSTM for SOH and RUL estimation of lithium-ion batteries[J]. Journal of Algorithms & Computational Technology, 2022, 16: 17483026221130598.",
        "Zhou D, Li Z, Zhu J, Zhang H, Hou L. State of Health Monitoring and Remaining Useful Life Prediction of Lithium-Ion Batteries Based on Temporal Convolutional Network[J]. IEEE Access, 2020, 8: 53307-53320.",
        "Ren L, Dong J, Wang X, et al. A Data-Driven Auto-CNN-LSTM Prediction Model for Lithium-Ion Battery Remaining Useful Life[J]. IEEE Transactions on Industrial Informatics, 2021, 17(5): 3478-3487.",
        "Wang Y, Zhu J, Cao L, et al. Long Short-Term Memory Network with Transfer Learning for Lithium-ion Battery Capacity Fade and Cycle Life Prediction[J]. Applied Energy, 2023, 346: 121660.",
    ]
    for ref in new_refs:
        append_reference_if_missing(doc, ref)

    try:
        doc.save(THESIS_PATH)
        print(THESIS_PATH)
    except PermissionError:
        doc.save(TEMP_OUTPUT_PATH)
        print(TEMP_OUTPUT_PATH)


if __name__ == "__main__":
    main()
